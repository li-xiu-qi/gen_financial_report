import hashlib
import json
import re
from os import getenv, makedirs, path
from pathlib import Path
from urllib.parse import quote, urlparse
import asyncio
import threading
from collections import defaultdict

import requests
import aiohttp
from dotenv import load_dotenv
from pyquery import PyQuery as pq

try:
    from fake_useragent import UserAgent
    ua = UserAgent()
except ImportError:
    ua = None

try:
    from zhipuai import ZhipuAI
except ImportError:
    ZhipuAI = None

from .utils.html2md import html2md
from .utils.clean_links import clean_markdown_links
from .utils.save_pdf_http import save_pdf
from .utils.remote_mineru_api_client import RemoteMineruAPIClient

load_dotenv()

# ====== 黑名单配置 ======
BLACKLIST_DOMAINS = {
    "zhihu.com",
    "zhuanlan.zhihu.com",
    "chinabgao.com",
}
# =======================

def _is_blacklisted(url: str) -> bool:
    """判断URL是否在黑名单域名内"""
    netloc = urlparse(url).netloc.lower()
    # 只要netloc包含黑名单域名即可
    for domain in BLACKLIST_DOMAINS:
        if domain in netloc:
            return True
    return False

# 域名锁管理器，确保同一域名只有一个请求在进行
class DomainLockManager:
    def __init__(self):
        self._locks = defaultdict(asyncio.Lock)
        self._main_lock = asyncio.Lock()
    
    async def get_lock(self, domain: str) -> asyncio.Lock:
        async with self._main_lock:
            return self._locks[domain]


def _ensure_cache_dir() -> str:
    cache_path = path.join(path.dirname(__file__), "cache")
    if not path.exists(cache_path):
        makedirs(cache_path)
    return cache_path


def md5_hash(text: str) -> str:
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def load_page_with_cache(
    url: str,
    cache_prefix: str,
    force_refresh: bool = False,
    cache_dir: str = None,
    search_api_url: str = None,
    pdf_base_url: str = None,
    timeout: int = 30,
) -> dict | None:
    # 黑名单处理
    if _is_blacklisted(url):
        print(f"黑名单域名，按超时处理: {url}")
        return None

    if cache_dir is None:
        cache_dir = _ensure_cache_dir()

    cache_file = path.join(cache_dir, f"{cache_prefix}.{md5_hash(url)}.json")
    
    print(f"\033[33mLoad page: {url}\033[0m")
    
    try:
        if force_refresh:
            raise FileNotFoundError
        with open(cache_file, "r", encoding="utf8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        if ua:
            user_agent = ua.random
        else:
            user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        
        headers = {
            "User-Agent": user_agent
        }
        
        cur_timeout = timeout
        
        page_info = None
        try:
            print(f"直接请求: {url}")
            direct_resp = requests.get(url, headers=headers, timeout=cur_timeout)
            if direct_resp.status_code == 200:
                page_info = {
                    "url": url,
                    "html": direct_resp.text
                }
                print(f"直接请求成功: {url}")
        except Exception as e:
            print(f"直接请求失败: {url}, {e}")
        
        if not page_info:
            search_api_url = search_api_url or getenv("SEARCH_URL")
            try:
                page_text = requests.get(
                    f"{search_api_url}/goto?query={quote(url)}", timeout=cur_timeout
                ).text
            except requests.Timeout:
                print(f"请求超时: {url} (timeout={cur_timeout}s)")
                return None
            except Exception as e:
                print(f"请求异常: {url}, {e}")
                return None
            if not page_text:
                return None
            page_info = json.loads(page_text)
        name, ext = path.splitext(page_info["url"])
        ext = re.sub(r"[^0-9a-zA-Z.]+", "", ext, flags=re.IGNORECASE)[0:10] or ".html"

        result_data = None

        allowed_exts = ["", ".html", ".htm", ".pdf", ".docx"]
        filter_exts = [
            ".xls", ".xlsx", ".csv", ".ppt", ".pptx", ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".svg", ".asp", ".txt", ".doc", ".rtf", ".odt", ".log", ".json", ".xml", ".js", ".css", ".zip", ".rar", ".7z", ".tar", ".gz", ".exe", ".apk", ".bin", ".dll", ".iso", ".mp3", ".mp4", ".avi", ".mov", ".wmv", ".flv", ".mkv", ".wav", ".wps"
        ]
        if ext.lower() in filter_exts:
            return None
        if ext == ".pdf":
            pdf_path = path.join(cache_dir, f'bing.{md5_hash(page_info["url"])}.pdf')
            save_pdf(page_info["url"], pdf_path)
            pdf_service_url = pdf_base_url or getenv("PDF_BASE_URL")
            pdf_client = RemoteMineruAPIClient(server_url=pdf_service_url)
            health = pdf_client.health_check()
            if health and health.get("status") == "ok":
                results = pdf_client.convert_pdf_to_md(file_paths=[pdf_path])
                if results and results[0].get("md_content"):
                    cleaned_md = clean_markdown_links(results[0]["md_content"])
                    print(f"PDF内容长度: {len(cleaned_md) // 1024} KB")
                    result_data = dict(
                        url=url,
                        data_source_type="pdf",
                        md=cleaned_md,
                    )
        elif ext == ".docx":
            docx_path = path.join(cache_dir, f'bing.{md5_hash(page_info["url"])}.docx')
            try:
                with requests.get(page_info["url"], stream=True, timeout=cur_timeout) as r:
                    r.raise_for_status()
                    with open(docx_path, "wb") as f:
                        for chunk in r.iter_content(chunk_size=8192):
                            f.write(chunk)
            except Exception as e:
                print(f"下载 docx 失败: {page_info['url']}, {e}")
                return None
            try:
                from docx import Document
                doc = Document(docx_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                docx_md = '\n\n'.join(paragraphs)
                cleaned_md = clean_markdown_links(docx_md)
                result_data = dict(
                    url=url,
                    data_source_type="docx",
                    md=cleaned_md,
                )
            except Exception as e:
                print(f"读取 docx 失败: {docx_path}, {e}")
                return None
        else:
            raw_md = html2md(page_info["html"])
            cleaned_md = clean_markdown_links(raw_md)
            result_data = dict(url=url, data_source_type="html", md=cleaned_md)

        if result_data:
            with open(cache_file, "w", encoding="utf8") as f:
                json.dump(result_data, f, ensure_ascii=False, indent=4)
        return result_data


async def load_page_with_cache_async(
    url: str,
    cache_prefix: str,
    force_refresh: bool = False,
    cache_dir: str = None,
    search_api_url: str = None,
    pdf_base_url: str = None,
    timeout: int = 30,
    domain_lock_manager: DomainLockManager = None,
) -> dict | None:
    # 黑名单处理
    if _is_blacklisted(url):
        print(f"黑名单域名，按超时处理: {url}")
        return None

    if cache_dir is None:
        cache_dir = _ensure_cache_dir()

    cache_file = path.join(cache_dir, f"{cache_prefix}.{md5_hash(url)}.json")
    
    print(f"\033[33mLoad page: {url}\033[0m")
    
    try:
        if not force_refresh:
            with open(cache_file, "r", encoding="utf8") as f:
                return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        pass
    
    domain = urlparse(url).netloc
    if domain_lock_manager:
        domain_lock = await domain_lock_manager.get_lock(domain)
        async with domain_lock:
            return await _fetch_page_content_async(
                url, cache_file, search_api_url, pdf_base_url, timeout, cache_dir
            )
    else:
        return await _fetch_page_content_async(
            url, cache_file, search_api_url, pdf_base_url, timeout, cache_dir
        )


async def _fetch_page_content_async(
    url: str,
    cache_file: str,
    search_api_url: str,
    pdf_base_url: str,
    timeout: int,
    cache_dir: str,
) -> dict | None:
    # 黑名单处理
    if _is_blacklisted(url):
        print(f"黑名单域名，按超时处理: {url}")
        return None

    if ua:
        user_agent = ua.random
    else:
        user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    
    headers = {
        "User-Agent": user_agent
    }
    
    cur_timeout = timeout
    page_info = None
    
    try:
        print(f"直接请求: {url}")
        timeout_obj = aiohttp.ClientTimeout(total=cur_timeout)
        async with aiohttp.ClientSession(timeout=timeout_obj, headers=headers) as session:
            async with session.get(url) as response:
                if response.status == 200:
                    html_content = await response.text()
                    page_info = {
                        "url": url,
                        "html": html_content
                    }
                    print(f"直接请求成功: {url}")
    except Exception as e:
        print(f"直接请求失败: {url}, {e}")
    
    if not page_info:
        search_api_url = search_api_url or getenv("SEARCH_URL")
        try:
            timeout_obj = aiohttp.ClientTimeout(total=cur_timeout)
            async with aiohttp.ClientSession(timeout=timeout_obj) as session:
                async with session.get(f"{search_api_url}/goto?query={quote(url)}") as response:
                    page_text = await response.text()
        except asyncio.TimeoutError:
            print(f"请求超时: {url} (timeout={cur_timeout}s)")
            return None
        except Exception as e:
            print(f"请求异常: {url}, {e}")
            return None
        if not page_text:
            return None
        page_info = json.loads(page_text)
    
    name, ext = path.splitext(page_info["url"])
    ext = re.sub(r"[^0-9a-zA-Z.]+", "", ext, flags=re.IGNORECASE)[0:10] or ".html"

    result_data = None
    
    filter_exts = [
        ".xls", ".xlsx", ".csv", ".ppt", ".pptx", ".jpg", ".jpeg", ".png", ".gif", 
        ".bmp", ".webp", ".svg", ".asp", ".txt", ".doc", ".rtf", ".odt", ".log", 
        ".json", ".xml", ".js", ".css", ".zip", ".rar", ".7z", ".tar", ".gz", 
        ".exe", ".apk", ".bin", ".dll", ".iso", ".mp3", ".mp4", ".avi", ".mov", 
        ".wmv", ".flv", ".mkv", ".wav", ".wps"
    ]
    if ext.lower() in filter_exts:
        return None
    
    if ext == ".pdf":
        pdf_path = path.join(cache_dir, f'bing.{md5_hash(page_info["url"])}.pdf')
        save_pdf(page_info["url"], pdf_path)
        pdf_service_url = pdf_base_url or getenv("PDF_BASE_URL")
        pdf_client = RemoteMineruAPIClient(server_url=pdf_service_url)
        health = pdf_client.health_check()
        if health and health.get("status") == "ok":
            results = pdf_client.convert_pdf_to_md(file_paths=[pdf_path])
            if results and results[0].get("md_content"):
                cleaned_md = clean_markdown_links(results[0]["md_content"])
                print(f"PDF内容长度: {len(cleaned_md) // 1024} KB")
                result_data = dict(
                    url=url,
                    data_source_type="pdf",
                    md=cleaned_md,
                )
    elif ext == ".docx":
        docx_path = path.join(cache_dir, f'bing.{md5_hash(page_info["url"])}.docx')
        try:
            timeout_obj = aiohttp.ClientTimeout(total=cur_timeout)
            async with aiohttp.ClientSession(timeout=timeout_obj) as session:
                async with session.get(page_info["url"]) as response:
                    response.raise_for_status()
                    with open(docx_path, "wb") as f:
                        async for chunk in response.content.iter_chunked(8192):
                            f.write(chunk)
        except Exception as e:
            print(f"下载 docx 失败: {page_info['url']}, {e}")
            return None
        
        try:
            from docx import Document
            doc = Document(docx_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            docx_md = '\n\n'.join(paragraphs)
            cleaned_md = clean_markdown_links(docx_md)
            result_data = dict(
                url=url,
                data_source_type="docx",
                md=cleaned_md,
            )
        except Exception as e:
            print(f"读取 docx 失败: {docx_path}, {e}")
            return None
    else:
        raw_md = html2md(page_info["html"])
        cleaned_md = clean_markdown_links(raw_md)
        result_data = dict(url=url, data_source_type="html", md=cleaned_md)

    if result_data:
        with open(cache_file, "w", encoding="utf8") as f:
            json.dump(result_data, f, ensure_ascii=False, indent=4)
    
    return result_data


def bing_search_with_cache(
    query: str,
    search_api_url: str,
    total: int = 50,
    cn: int = None,
    force_refresh: bool = False,
    cache_dir: str = None,
    searched_urls: set = None,
    pdf_base_url: str = None,
    timeout: int = 30,
    use_async: bool = True,
    rate_limit_delay: float = 0.5,
) -> list:
    if use_async:
        return _bing_search_async(
            query, search_api_url, total, cn, force_refresh, 
            cache_dir, searched_urls, pdf_base_url, timeout, rate_limit_delay
        )
    else:
        return _bing_search_sync(
            query, search_api_url, total, cn, force_refresh, 
            cache_dir, searched_urls, pdf_base_url, timeout, rate_limit_delay
        )


def _bing_search_sync(
    query: str,
    search_api_url: str,
    total: int,
    cn: int,
    force_refresh: bool,
    cache_dir: str,
    searched_urls: set,
    pdf_base_url: str,
    timeout: int,
    rate_limit_delay: float,
) -> list:
    if cache_dir is None:
        cache_dir = _ensure_cache_dir()

    search_api_url = search_api_url or getenv("SEARCH_URL")
    search_url = f"{search_api_url}/bing?query={quote(query)}&total={total or 10}"
    if cn:
        search_url += "&cn=1"
    else:
        search_url += "&cn=0"
    search_results = requests.get(search_url).json()

    results = []
    if searched_urls is None:
        searched_urls = set()

    timeout_count = 0

    for i, item in enumerate(search_results):
        url = item["url"]
        if url in searched_urls:
            print(f"跳过已抓取: {url}")
            continue

        if i > 0 and rate_limit_delay > 0:
            print(f"⏳ 等待 {rate_limit_delay} 秒，避免请求过于频繁...")
            import time
            time.sleep(rate_limit_delay)

        res = load_page_with_cache(
            url,
            cache_prefix="bing",
            force_refresh=force_refresh,
            cache_dir=cache_dir,
            search_api_url=search_api_url,
            pdf_base_url=pdf_base_url,
            timeout=timeout,
        )
        if not res:
            timeout_count += 1
            print(f"请求超时: {url} (timeout={timeout}s)")
            continue
        results.append(
            dict(
                title=item["title"],
                **res,
            )
        )
        searched_urls.add(url)
    return results


def _bing_search_async(
    query: str,
    search_api_url: str,
    total: int,
    cn: int,
    force_refresh: bool,
    cache_dir: str,
    searched_urls: set,
    pdf_base_url: str,
    timeout: int,
    rate_limit_delay: float,
) -> list:
    try:
        loop = asyncio.get_running_loop()
        def run_in_thread():
            return asyncio.run(_async_bing_search_impl(
                query, search_api_url, total, cn, force_refresh,
                cache_dir, searched_urls, pdf_base_url, timeout, rate_limit_delay
            ))
        
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(run_in_thread)
            return future.result()
    except RuntimeError:
        return asyncio.run(_async_bing_search_impl(
            query, search_api_url, total, cn, force_refresh,
            cache_dir, searched_urls, pdf_base_url, timeout, rate_limit_delay
        ))


async def _async_bing_search_impl(
    query: str,
    search_api_url: str,
    total: int,
    cn: int,
    force_refresh: bool,
    cache_dir: str,
    searched_urls: set,
    pdf_base_url: str,
    timeout: int,
    rate_limit_delay: float,
) -> list:
    if cache_dir is None:
        cache_dir = _ensure_cache_dir()

    search_api_url = search_api_url or getenv("SEARCH_URL")
    search_url = f"{search_api_url}/bing?query={quote(query)}&total={total or 10}"
    if cn:
        search_url += "&cn=1"
    else:
        search_url += "&cn=0"
    
    search_results = requests.get(search_url).json()
    
    results = []
    if searched_urls is None:
        searched_urls = set()
    
    domain_lock_manager = DomainLockManager()
    tasks = []
    
    async def fetch_single_item(i: int, item: dict):
        url = item["url"]
        if url in searched_urls:
            print(f"跳过已抓取: {url}")
            return None
        
        print(f"\033[33mProcessing [{i+1}/{len(search_results)}]: {url}\033[0m")
        
        res = await load_page_with_cache_async(
            url,
            cache_prefix="bing",
            force_refresh=force_refresh,
            cache_dir=cache_dir,
            search_api_url=search_api_url,
            pdf_base_url=pdf_base_url,
            timeout=timeout,
            domain_lock_manager=domain_lock_manager,
        )
        
        if not res:
            print(f"请求失败: {url}")
            return None
        
        searched_urls.add(url)
        return dict(
            title=item["title"],
            **res,
        )
    
    for i, item in enumerate(search_results):
        task = fetch_single_item(i, item)
        tasks.append(task)
    
    task_results = await asyncio.gather(*tasks, return_exceptions=True)
    
    for result in task_results:
        if result is not None and not isinstance(result, Exception):
            results.append(result)
        elif isinstance(result, Exception):
            print(f"任务执行异常: {result}")
    
    return results


def get_tonghuashun_data(
    tonghuashun_total_code: str,
    search_api_url: str,
    force_refresh: bool = False,
    cache_dir: str = None,
    pdf_base_url: str = None,
    timeout: int = 30,
) -> dict:
    if cache_dir is None:
        cache_dir = _ensure_cache_dir()

    search_api_url = search_api_url or getenv("SEARCH_URL")
    normalized_code = tonghuashun_total_code

    url = f"{search_api_url}/10jqka?query={normalized_code}"
    print(f"Fetching data for {normalized_code} from {url}")
    import time
    from http.client import IncompleteRead
    from requests.exceptions import ChunkedEncodingError, ConnectionError
    max_retries = 3
    company_data = None
    for attempt in range(max_retries):
        try:
            resp = requests.get(url, timeout=timeout)
            company_data = resp.json()
            break
        except requests.Timeout:
            print(f"同花顺接口超时: {url} (timeout={timeout}s)，重试 {attempt+1}/{max_retries}")
            time.sleep(2)
        except (IncompleteRead, ChunkedEncodingError, ConnectionError) as e:
            print(f"同花顺接口连接异常: {url}, {e}，重试 {attempt+1}/{max_retries}")
            time.sleep(2)
        except Exception as e:
            print(f"同花顺接口其它异常: {url}, {e}")
            break
    else:
        print(f"同花顺接口多次重试后仍失败: {url}")
        return dict(navs=[], news=[])

    result = dict(navs=[], news=[])

    for nav_item in company_data:
        html_content = nav_item.pop("html", "")
        if not html_content:
            continue

        if "新闻" not in nav_item["title"] and "公告" not in nav_item["title"]:
            raw_md = html2md(html_content)
            cleaned_md = clean_markdown_links(raw_md)
            nav_item.update(
                dict(
                    md=cleaned_md,
                    data_source_type="html",
                )
            )
            result["navs"].append(nav_item)
        else:
            doc = pq(html_content)
            for link_element in doc.find("dl dt a, td a[tag=reports]"):
                link = pq(link_element)
                url, title = link.attr("href"), link.attr("title")

                if not url or not title:
                    continue

                print(f"\033[34mFetching news: {title} ({url})\033[0m")

                res = load_page_with_cache(
                    url,
                    cache_prefix="10jqka",
                    force_refresh=force_refresh,
                    cache_dir=cache_dir,
                    search_api_url=search_api_url,
                    pdf_base_url=pdf_base_url,
                    timeout=timeout,
                )

                if res:
                    if "md" in res:
                        res["md"] = clean_markdown_links(res["md"])
                    res.update(
                        dict(
                            title=title,
                            url=url,
                        )
                    )
                    result["news"].append(res)

    return result


def zhipu_search_with_cache(
    query: str,
    count: int = 50,
    search_engine: str = "Search-Pro",
    force_refresh: bool = False,
    cache_dir: str = None,
    searched_urls: set = None,
    pdf_base_url: str = None,
    timeout: int = 30,
    use_async: bool = True,
    zhipu_api_key: str = None,
    rate_limit_delay: float = 0.5,
) -> list:
    if ZhipuAI is None:
        print("智谱AI库未安装，请先安装: pip install zhipuai")
        return []
    
    if use_async:
        return _zhipu_search_async(
            query, count, search_engine, force_refresh, 
            cache_dir, searched_urls, pdf_base_url, timeout, zhipu_api_key, rate_limit_delay
        )
    else:
        return _zhipu_search_sync(
            query, count, search_engine, force_refresh, 
            cache_dir, searched_urls, pdf_base_url, timeout, zhipu_api_key, rate_limit_delay
        )


def _zhipu_search_sync(
    query: str,
    count: int,
    search_engine: str,
    force_refresh: bool,
    cache_dir: str,
    searched_urls: set,
    pdf_base_url: str,
    timeout: int,
    zhipu_api_key: str,
    rate_limit_delay: float,
) -> list:
    if cache_dir is None:
        cache_dir = _ensure_cache_dir()

    api_key = zhipu_api_key or getenv("ZHIPU_API_KEY")
    if not api_key:
        print("未找到智谱AI API密钥，请设置ZHIPU_API_KEY环境变量或传入zhipu_api_key参数")
        return []

    try:
        client = ZhipuAI(api_key=api_key)
        
        print(f"\033[33m智谱搜索: {query}\033[0m")
        
        response = client.web_search.web_search(
            search_engine=search_engine,
            search_query=query,
            count=count,
        )
        
        results = []
        if searched_urls is None:
            searched_urls = set()

        for idx, result in enumerate(response.search_result, 1):
            title = getattr(result, 'title', '')
            url = getattr(result, 'link', '')
            content = getattr(result, 'content', '')
            
            if not url:
                continue
                
            if url in searched_urls:
                print(f"跳过已抓取: {url}")
                continue

            print(f"\033[33m处理第{idx}条结果: {title}\033[0m")

            res = load_page_with_cache(
                url,
                cache_prefix="zhipu",
                force_refresh=force_refresh,
                cache_dir=cache_dir,
                pdf_base_url=pdf_base_url,
                timeout=timeout,
            )
            
            if res:
                results.append(
                    dict(
                        title=title,
                        summary=content,
                        **res,
                    )
                )
            else:
                results.append(
                    dict(
                        title=title,
                        url=url,
                        md=content,
                        data_source_type="search_summary",
                        summary=content,
                    )
                )
            
            searched_urls.add(url)
            
    except Exception as e:
        print(f"智谱搜索异常: {e}")
        return []

    return results


def _zhipu_search_async(
    query: str,
    count: int,
    search_engine: str,
    force_refresh: bool,
    cache_dir: str,
    searched_urls: set,
    pdf_base_url: str,
    timeout: int,
    zhipu_api_key: str,
    rate_limit_delay: float,
) -> list:
    try:
        loop = asyncio.get_running_loop()
        def run_in_thread():
            return asyncio.run(_async_zhipu_search_impl(
                query, count, search_engine, force_refresh,
                cache_dir, searched_urls, pdf_base_url, timeout, zhipu_api_key, rate_limit_delay
            ))
        
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(run_in_thread)
            return future.result()
    except RuntimeError:
        return asyncio.run(_async_zhipu_search_impl(
            query, count, search_engine, force_refresh,
            cache_dir, searched_urls, pdf_base_url, timeout, zhipu_api_key, rate_limit_delay
        ))


async def _async_zhipu_search_impl(
    query: str,
    count: int,
    search_engine: str,
    force_refresh: bool,
    cache_dir: str,
    searched_urls: set,
    pdf_base_url: str,
    timeout: int,
    zhipu_api_key: str,
    rate_limit_delay: float,
) -> list:
    if cache_dir is None:
        cache_dir = _ensure_cache_dir()

    api_key = zhipu_api_key or getenv("ZHIPU_API_KEY")
    if not api_key:
        print("未找到智谱AI API密钥，请设置ZHIPU_API_KEY环境变量或传入zhipu_api_key参数")
        return []

    try:
        client = ZhipuAI(api_key=api_key)
        
        print(f"\033[33m智谱搜索: {query}\033[0m")
        
        response = client.web_search.web_search(
            search_engine=search_engine,
            search_query=query,
            count=count,
        )
        
        results = []
        if searched_urls is None:
            searched_urls = set()
        
        domain_lock_manager = DomainLockManager()
        tasks = []
        
        async def fetch_single_item(idx: int, result_item):
            title = getattr(result_item, 'title', '')
            url = getattr(result_item, 'link', '')
            content = getattr(result_item, 'content', '')
            
            if not url:
                return None
                
            if url in searched_urls:
                print(f"跳过已抓取: {url}")
                return None
            
            print(f"\033[33m处理第{idx}条结果: {title} - {url}\033[0m")
            
            res = await load_page_with_cache_async(
                url,
                cache_prefix="zhipu",
                force_refresh=force_refresh,
                cache_dir=cache_dir,
                pdf_base_url=pdf_base_url,
                timeout=timeout,
                domain_lock_manager=domain_lock_manager,
            )
            
            searched_urls.add(url)
            
            if res:
                return dict(
                    title=title,
                    summary=content,
                    **res,
                )
            else:
                return dict(
                    title=title,
                    url=url,
                    md=content,
                    data_source_type="search_summary",
                    summary=content,
                )
        
        for idx, result_item in enumerate(response.search_result, 1):
            task = fetch_single_item(idx, result_item)
            tasks.append(task)
        
        task_results = await asyncio.gather(*tasks, return_exceptions=True)
        
        for result in task_results:
            if result is not None and not isinstance(result, Exception):
                results.append(result)
            elif isinstance(result, Exception):
                print(f"任务执行异常: {result}")
                
    except Exception as e:
        print(f"智谱搜索异常: {e}")
        return []
    
    return results
